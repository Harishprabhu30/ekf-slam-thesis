#!/usr/bin/env python3

from pathlib import Path
import math
import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


BASE_DIR = Path("/media/vgtu/New Volume1/Harish_Thesis/ros2_ws/analysis/results_v6_lighting_sweep")
METRICS_CSV = BASE_DIR / "metrics" / "v6_all_trial_metrics.csv"
OUT_DIR = BASE_DIR / "thesis_chapter_4_tables"

REP_TRIAL = "t03"
PHASE_LIGHTING = "lowlight"

LIGHT_ORDER = ["bright", "dim", "lowlight"]
METHOD_ORDER = ["wheel", "ekf", "orbslam3"]

LIGHT_LABEL = {
    "bright": "Bright",
    "dim": "Dim",
    "lowlight": "Low-light",
    "low-light": "Low-light",
}

METHOD_LABEL = {
    "wheel": "Wheel",
    "ekf": "EKF",
    "orbslam3": "ORB-SLAM3",
    "orb-slam3": "ORB-SLAM3",
}


ALIASES = {
    "lighting": ["lighting", "light", "condition"],
    "trial": ["trial", "run", "trial_id"],
    "method": ["method", "estimator", "config"],
    "ate_rmse": ["ate_rmse_m", "ate_rmse", "ATE RMSE [m]", "ATE RMSE"],
    "ate_mean": ["ate_mean_m", "ate_mean", "ATE mean", "ATE Mean"],
    "ate_max": ["ate_max_m", "ate_max", "ATE max", "ATE Max"],
    "rpe_trans": [
        "rpe_trans_rmse_m_1.0s",
        "rpe_trans_rmse_m_1s",
        "rpe_trans_rmse_m",
        "rpe_translation_rmse_m",
        "RPE_1s [m]",
        "RPE trans RMSE",
    ],
    "yaw": ["yaw_abs_mean_deg", "yaw_error_deg", "mean_yaw_error_deg", "Yaw [deg]"],
    "drift": ["drift_rate_m_per_m", "drift_rate", "Drift [m/m]"],
    "pose_ratio": ["pose_count_ratio", "pose_ratio", "Pose ratio"],
    "norm_ate": ["ate_norm_rmse", "normalized_ate", "norm_ate", "nATE"],
}


def find_col(df, key, required=True):
    for c in ALIASES[key]:
        if c in df.columns:
            return c
    if required:
        raise KeyError(f"Could not find column for '{key}'. Available columns: {list(df.columns)}")
    return None


def normalize_light(x):
    x = str(x).strip().lower().replace("-", "")
    if x == "lowlight":
        return "lowlight"
    return x


def normalize_method(x):
    x = str(x).strip().lower()
    if "orb" in x:
        return "orbslam3"
    if "ekf" in x:
        return "ekf"
    if "wheel" in x:
        return "wheel"
    return x


def fmt_mean_std(mean_val, std_val, decimals=3):
    if pd.isna(mean_val):
        return "N/A"
    if pd.isna(std_val):
        std_val = 0.0
    return f"{mean_val:.{decimals}f} ± {std_val:.{decimals}f}"


def load_metrics():
    if not METRICS_CSV.exists():
        raise FileNotFoundError(f"Metrics file not found:\n{METRICS_CSV}")

    df = pd.read_csv(METRICS_CSV)

    col_light = find_col(df, "lighting")
    col_trial = find_col(df, "trial")
    col_method = find_col(df, "method")

    rename_map = {
        col_light: "lighting",
        col_trial: "trial",
        col_method: "method",
        find_col(df, "ate_rmse"): "ate_rmse_m",
        find_col(df, "rpe_trans"): "rpe_trans_rmse_m",
        find_col(df, "yaw"): "yaw_abs_mean_deg",
        find_col(df, "drift"): "drift_rate_m_per_m",
        find_col(df, "pose_ratio"): "pose_count_ratio",
    }

    optional_cols = {
        "ate_mean": "ate_mean_m",
        "ate_max": "ate_max_m",
        "norm_ate": "ate_norm_rmse",
    }

    for key, new_name in optional_cols.items():
        col = find_col(df, key, required=False)
        if col is not None:
            rename_map[col] = new_name

    df = df.rename(columns=rename_map)

    df["lighting"] = df["lighting"].apply(normalize_light)
    df["method"] = df["method"].apply(normalize_method)
    df["trial"] = df["trial"].astype(str).str.strip()

    return df


def make_table_4_1(df):
    rows = []

    metrics = [
        ("ate_rmse_m", "ATE RMSE (m)", 3),
        ("rpe_trans_rmse_m", "RPE trans. RMSE (m)", 3),
        ("yaw_abs_mean_deg", "Mean yaw error (deg)", 2),
        ("drift_rate_m_per_m", "Drift rate (m/m)", 3),
        ("pose_count_ratio", "Pose count ratio", 2),
    ]

    for light in LIGHT_ORDER:
        for method in METHOD_ORDER:
            sub = df[(df["lighting"] == light) & (df["method"] == method)]
            row = {
                "Lighting": LIGHT_LABEL[light],
                "Method": METHOD_LABEL[method],
            }

            for col, out_name, decimals in metrics:
                row[out_name] = fmt_mean_std(sub[col].mean(), sub[col].std(ddof=1), decimals)

            rows.append(row)

    return pd.DataFrame(rows)


def make_table_4_2(df):
    rows = []

    for light in LIGHT_ORDER:
        for method in METHOD_ORDER:
            sub = df[(df["lighting"] == light) & (df["method"] == method) & (df["trial"] == REP_TRIAL)]
            if sub.empty:
                trial = "N/A"
                criterion = f"{REP_TRIAL} not found in metrics file"
            else:
                trial = REP_TRIAL
                criterion = "Fixed representative trial used for Chapter 4 visualizations"

            rows.append({
                "Lighting condition": LIGHT_LABEL[light],
                "Method": METHOD_LABEL[method],
                "Representative trial": trial,
                "Selection criterion": criterion,
            })

    return pd.DataFrame(rows)


def find_time_col(df):
    for c in ["t", "time", "time_s", "timestamp", "stamp", "sec"]:
        if c in df.columns:
            return c
    return None


def find_ate_col(df):
    for c in ["ate_m", "ate", "ate_error_m", "position_error_m"]:
        if c in df.columns:
            return c
    return None


def find_xy_pair(df, prefix_candidates):
    cols = list(df.columns)

    for px in prefix_candidates:
        candidates = [
            (f"{px}_x", f"{px}_y"),
            (f"x_{px}", f"y_{px}"),
            (f"{px}x", f"{px}y"),
            (f"{px}_aligned_x", f"{px}_aligned_y"),
            (f"{px}_x_aligned", f"{px}_y_aligned"),
        ]
        for x_col, y_col in candidates:
            if x_col in cols and y_col in cols:
                return x_col, y_col

    return None, None


def compute_ate_series(aligned_df):
    ate_col = find_ate_col(aligned_df)
    if ate_col:
        return aligned_df[ate_col].astype(float)

    gt_x, gt_y = find_xy_pair(aligned_df, ["gt", "groundtruth", "ground_truth"])
    est_x, est_y = find_xy_pair(aligned_df, ["est", "aligned", "estimate", "traj", "pose"])

    if gt_x is None or est_x is None:
        return None

    return ((aligned_df[est_x] - aligned_df[gt_x]) ** 2 + (aligned_df[est_y] - aligned_df[gt_y]) ** 2) ** 0.5


def load_phase_intervals():
    phase_file = BASE_DIR / "extracted" / f"{PHASE_LIGHTING}_{REP_TRIAL}_wheel_phase_events.csv"
    if not phase_file.exists():
        return None

    ph = pd.read_csv(phase_file)
    time_col = find_time_col(ph)

    phase_col = None
    for c in ["phase", "name", "event", "label", "data"]:
        if c in ph.columns:
            phase_col = c
            break

    if time_col is None or phase_col is None:
        return None

    ph = ph[[time_col, phase_col]].dropna().sort_values(time_col).reset_index(drop=True)
    if ph.empty:
        return None

    intervals = []
    for i in range(len(ph)):
        start = float(ph.loc[i, time_col])
        end = float(ph.loc[i + 1, time_col]) if i + 1 < len(ph) else math.inf
        raw_name = str(ph.loc[i, phase_col]).lower()
        canonical = canonical_phase(raw_name)
        if canonical:
            intervals.append((canonical, start, end))

    return intervals


def canonical_phase(name):
    name = name.lower()

    if "rest" in name or "stop" in name or "idle" in name:
        return "Rest / stop"
    if "straight" in name:
        return "Straight motion"
    if "square" in name or "turn" in name:
        return "Square turns"
    if "rotation" in name or "rotate" in name or "clockwise" in name:
        return "Rotation"
    if "curved" in name or "curve" in name or "circle" in name or "circular" in name:
        return "Curved motion"

    return None


def make_table_4_3():
    phase_order = ["Rest / stop", "Straight motion", "Square turns", "Rotation", "Curved motion"]
    observations = {
        "Rest / stop": "Stability/noise",
        "Straight motion": "Translational drift",
        "Square turns": "Turning/yaw error",
        "Rotation": "Orientation difficulty",
        "Curved motion": "Combined motion error",
    }

    intervals = load_phase_intervals()
    rows = []

    if intervals is None:
        for phase in phase_order:
            rows.append({
                "Phase": phase,
                "Wheel ATE RMSE": "N/A",
                "EKF ATE RMSE": "N/A",
                "ORB-SLAM3 ATE RMSE": "N/A",
                "Main observation": observations[phase],
            })
        return pd.DataFrame(rows)

    method_phase_values = {m: {p: [] for p in phase_order} for m in METHOD_ORDER}

    for method in METHOD_ORDER:
        aligned_file = BASE_DIR / "aligned" / f"{PHASE_LIGHTING}_{REP_TRIAL}_{method}_aligned.csv"
        if not aligned_file.exists():
            continue

        adf = pd.read_csv(aligned_file)
        time_col = find_time_col(adf)
        if time_col is None:
            continue

        ate = compute_ate_series(adf)
        if ate is None:
            continue

        adf = adf.copy()
        adf["_ate"] = ate.astype(float)

        for phase, start, end in intervals:
            sub = adf[(adf[time_col] >= start) & (adf[time_col] < end)]
            if len(sub) > 0:
                rmse = math.sqrt((sub["_ate"] ** 2).mean())
                method_phase_values[method][phase].append(rmse)

    for phase in phase_order:
        row = {
            "Phase": phase,
            "Main observation": observations[phase],
        }

        for method, label in [("wheel", "Wheel ATE RMSE"), ("ekf", "EKF ATE RMSE"), ("orbslam3", "ORB-SLAM3 ATE RMSE")]:
            vals = method_phase_values[method][phase]
            row[label] = f"{sum(vals) / len(vals):.3f}" if vals else "N/A"

        rows.append(row)

    return pd.DataFrame(rows)[
        ["Phase", "Wheel ATE RMSE", "EKF ATE RMSE", "ORB-SLAM3 ATE RMSE", "Main observation"]
    ]


def make_table_4_4(df):
    rows = []
    observations = {
        "bright": "Feature tracking under bright illumination",
        "dim": "Reduced visual robustness",
        "lowlight": "High trial dependence and feature-extraction sensitivity",
    }

    for light in LIGHT_ORDER:
        sub = df[(df["lighting"] == light) & (df["method"] == "orbslam3")]

        rows.append({
            "Lighting condition": LIGHT_LABEL[light],
            "ATE RMSE": fmt_mean_std(sub["ate_rmse_m"].mean(), sub["ate_rmse_m"].std(ddof=1), 3),
            "Pose count ratio": fmt_mean_std(sub["pose_count_ratio"].mean(), sub["pose_count_ratio"].std(ddof=1), 2),
            "RPE translation RMSE": fmt_mean_std(sub["rpe_trans_rmse_m"].mean(), sub["rpe_trans_rmse_m"].std(ddof=1), 3),
            "Main observation": observations[light],
        })

    return pd.DataFrame(rows)


def make_table_c_1(df):
    cols = [
        "lighting", "trial", "method",
        "ate_rmse_m", "ate_mean_m", "ate_max_m",
        "rpe_trans_rmse_m", "yaw_abs_mean_deg",
        "drift_rate_m_per_m", "pose_count_ratio"
    ]

    for c in ["ate_mean_m", "ate_max_m"]:
        if c not in df.columns:
            df[c] = pd.NA

    out = df[cols].copy()

    out["lighting"] = out["lighting"].map(LIGHT_LABEL)
    out["method"] = out["method"].map(METHOD_LABEL)

    out = out.rename(columns={
        "lighting": "Lighting",
        "trial": "Trial",
        "method": "Method",
        "ate_rmse_m": "ATE RMSE",
        "ate_mean_m": "ATE mean",
        "ate_max_m": "ATE max",
        "rpe_trans_rmse_m": "RPE trans RMSE",
        "yaw_abs_mean_deg": "Yaw error",
        "drift_rate_m_per_m": "Drift rate",
        "pose_count_ratio": "Pose count ratio",
    })

    numeric_cols = [
        "ATE RMSE", "ATE mean", "ATE max",
        "RPE trans RMSE", "Yaw error",
        "Drift rate", "Pose count ratio"
    ]
    for c in numeric_cols:
        out[c] = pd.to_numeric(out[c], errors="coerce").round(3)

    out["Lighting_sort"] = out["Lighting"].map({"Bright": 0, "Dim": 1, "Low-light": 2})
    out["Method_sort"] = out["Method"].map({"Wheel": 0, "EKF": 1, "ORB-SLAM3": 2})
    out = out.sort_values(["Lighting_sort", "Trial", "Method_sort"]).drop(columns=["Lighting_sort", "Method_sort"])

    return out


def save_csvs(tables):
    for name, df in tables.items():
        path = OUT_DIR / f"{name}.csv"
        df.to_csv(path, index=False)
        print(f"[SAVED] {path}")


def save_xlsx(tables):
    path = OUT_DIR / "thesis_chapter_4_tables.xlsx"

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for name, df in tables.items():
            sheet = name.replace("table_", "T").replace("_", " ")[:31]
            df.to_excel(writer, sheet_name=sheet, index=False)

            ws = writer.book[sheet]
            ws.freeze_panes = "A2"

            for col_cells in ws.columns:
                max_len = 0
                col_letter = col_cells[0].column_letter
                for cell in col_cells:
                    max_len = max(max_len, len(str(cell.value)) if cell.value is not None else 0)
                ws.column_dimensions[col_letter].width = min(max_len + 2, 45)

    print(f"[SAVED] {path}")


def add_docx_table(doc, caption, df):
    doc.add_paragraph(caption)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if pd.isna(value) else str(value)

    doc.add_paragraph("")


def save_docx(tables):
    if not DOCX_AVAILABLE:
        print("[WARNING] python-docx is not installed. DOCX export skipped.")
        print("Install with: pip3 install python-docx")
        return

    doc = Document()
    doc.add_heading("Chapter 4 Tables", level=1)

    captions = {
        "table_4_1_main_quantitative_result_summary":
            "Table 4.1. Mean and standard deviation of localization metrics across three trials for each lighting condition and estimator.",
        "table_4_2_representative_trial_selection":
            "Table 4.2. Representative trials selected for detailed trajectory and time-series visualization.",
        "table_4_3_phase_based_error_summary_lowlight_t03":
            "Table 4.3. Phase-based localization error summary showing estimator performance during different motion segments.",
        "table_4_4_orbslam3_lighting_robustness_summary":
            "Table 4.4. ORB-SLAM3 robustness summary under different lighting conditions.",
        "table_C_1_full_per_trial_metrics":
            "Table C.1. Full per-trial localization metrics for all evaluated runs.",
    }

    for name, df in tables.items():
        add_docx_table(doc, captions.get(name, name), df)

    path = OUT_DIR / "thesis_chapter_4_tables.docx"
    doc.save(path)
    print(f"[SAVED] {path}")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    df = load_metrics()

    tables = {
        "table_4_1_main_quantitative_result_summary": make_table_4_1(df),
        "table_4_2_representative_trial_selection": make_table_4_2(df),
        "table_4_3_phase_based_error_summary_lowlight_t03": make_table_4_3(),
        "table_4_4_orbslam3_lighting_robustness_summary": make_table_4_4(df),
        "table_C_1_full_per_trial_metrics": make_table_c_1(df),
    }

    save_csvs(tables)
    save_xlsx(tables)
    save_docx(tables)

    print("\n[DONE] Thesis Chapter 4 tables exported to:")
    print(OUT_DIR)


if __name__ == "__main__":
    main()
